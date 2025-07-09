import tkinter as tk
from tkinter import filedialog, messagebox
import nbformat
from nbconvert import HTMLExporter
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import base64
import os

def convertir_ipynb_a_docx_con_imagenes(ruta_ipynb):
    try:
        # Leer el notebook
        with open(ruta_ipynb, 'r', encoding='utf-8') as f:
            notebook = nbformat.read(f, as_version=4)

        # Convertir a HTML con nbconvert
        exporter = HTMLExporter()
        exporter.template_name = 'classic'
        body, resources = exporter.from_notebook_node(notebook)

        # Analizar HTML con BeautifulSoup
        soup = BeautifulSoup(body, 'html.parser')
        doc = Document()

        for tag in soup.find_all(['h1', 'h2', 'h3', 'p', 'img', 'pre']):
            if tag.name == 'h1':
                doc.add_heading(tag.text.strip(), level=1)
            elif tag.name == 'h2':
                doc.add_heading(tag.text.strip(), level=2)
            elif tag.name == 'h3':
                doc.add_heading(tag.text.strip(), level=3)
            elif tag.name == 'p':
                doc.add_paragraph(tag.text.strip())
            elif tag.name == 'pre':
                # Bloque de código con fuente monoespaciada
                p = doc.add_paragraph()
                run = p.add_run(tag.text.strip())
                font = run.font
                font.name = 'Courier New'
                font.size = Pt(10)
                font.bold = False
                font.italic = False
                # Asegura Courier New en Word (compatibilidad)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
            elif tag.name == 'img' and 'src' in tag.attrs:
                img_data = tag['src']
                if img_data.startswith("data:image/png;base64,"):
                    base64_img = img_data.split(",")[1]
                    img_bytes = base64.b64decode(base64_img)
                    temp_img_path = "temp_output_image.png"
                    with open(temp_img_path, "wb") as img_file:
                        img_file.write(img_bytes)
                    doc.add_picture(temp_img_path, width=Inches(5.5))
                    os.remove(temp_img_path)

        # Guardar el documento
        carpeta = os.path.dirname(ruta_ipynb)
        nombre_base = os.path.splitext(os.path.basename(ruta_ipynb))[0]
        ruta_docx = os.path.join(carpeta, f"{nombre_base}.docx")
        doc.save(ruta_docx)

        messagebox.showinfo("Conversión exitosa", f"Archivo Word creado en:\n{ruta_docx}")
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo convertir el archivo:\n{e}")

def seleccionar_archivo():
    ruta = filedialog.askopenfilename(filetypes=[("Jupyter Notebook", "*.ipynb")])
    if ruta:
        entrada_ruta.delete(0, tk.END)
        entrada_ruta.insert(0, ruta)

def ejecutar_conversion():
    ruta = entrada_ruta.get()
    if ruta.endswith('.ipynb'):
        convertir_ipynb_a_docx_con_imagenes(ruta)
    else:
        messagebox.showwarning("Archivo inválido", "Selecciona un archivo .ipynb válido.")

# GUI
ventana = tk.Tk()
ventana.title("Convertir ipynb a Word (con gráficos y código)")
ventana.geometry("500x160")

tk.Label(ventana, text="Selecciona un archivo .ipynb").pack(pady=5)
entrada_ruta = tk.Entry(ventana, width=60)
entrada_ruta.pack(pady=5)

tk.Button(ventana, text="Buscar archivo", command=seleccionar_archivo).pack(pady=5)
tk.Button(ventana, text="Convertir a Word", command=ejecutar_conversion).pack(pady=5)

ventana.mainloop()
